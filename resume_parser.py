from docx import  Document

COMMON_SKILLS = [
    'python' , 'java' , 'django' , 'sql',"Traine" ,"internship","systems"
]

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:  # Avoid empty cells
                    full_text.append(cell_text.lower())

    return '\n'.join(full_text)

def extract_skills(text):
    found_skills = []
    for skill in COMMON_SKILLS:
        if skill.lower() in text:
            found_skills.append(skill)
    return found_skills
   